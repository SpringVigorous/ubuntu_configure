from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math

def ocr_results_to_word(ocr_results, output_path, img_width, img_height, doc_width=6):
    """
    将OCR识别结果按包围框位置输出到Word文档
    :param ocr_results: OCR结果列表，每个元素为[(x1,y1), (x2,y2), (x3,y3), (x4,y4), [text, score]]
    :param output_path: 输出Word路径（如"result.docx"）
    :param img_width: 原图宽度（像素）
    :param img_height: 原图高度（像素）
    :param doc_width: Word页面有效宽度（英寸，默认6英寸）
    """
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距，确保有效宽度
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # 计算像素到英寸的缩放比例（假设Word页面有效宽度为doc_width英寸）
    scale = doc_width / img_width  # 1像素对应的英寸数
    
    # 按y坐标排序（确保从上到下输出）
    ocr_results.sort(key=lambda x: min(y for (_, y) in x[0]))
    
    prev_y = 0  # 记录上一段落的y坐标，用于计算行距
    
    for bbox, (text, _) in ocr_results:
        # 提取包围框坐标
        (x1, y1), (x2, y2), (x3, y3), (x4, y4) = bbox
        
        # 计算文本位置（取左上角坐标作为参考点）
        text_x = min(x1, x4)  # 左边界x坐标
        text_y = min(y1, y2)  # 上边界y坐标
        
        # 计算与上一段落的垂直距离（转换为英寸）
        vertical_distance = (text_y - prev_y) * scale
        if vertical_distance < 0:
            vertical_distance = 0  # 防止重叠文本出现负距离
        
        # 添加段落
        para = doc.add_paragraph()
        
        # 设置段落前间距（控制垂直位置）
        para.paragraph_format.space_before = Pt(vertical_distance * 72)  # 1英寸=72磅
        
        # 设置左缩进（控制水平位置）
        para.paragraph_format.left_indent = Inches(text_x * scale)
        
        # 添加文本并设置字体（可选）
        run = para.add_run(text)
        run.font.size = Pt(10)  # 可根据实际情况调整字体大小
        
        # 更新上一段落y坐标（取当前文本框底部y坐标）
        prev_y = max(y3, y4)
    
    # 保存文档
    doc.save(output_path)
    print(f"Word文档已生成：{output_path}")

# 示例用法
if __name__ == "__main__":
    # 模拟OCR识别结果：[(包围框坐标), (文本, 置信度)]
    # 实际使用时替换为你的OCR结果
    ocr_results = [
        [
            [(50, 30), (200, 30), (200, 60), (50, 60)],  # 包围框
            ["标题文本", 0.98]  # 文本及置信度
        ],
        [
            [(50, 80), (400, 80), (400, 110), (50, 110)],
            ["这是第一行内容，位于标题下方", 0.95]
        ],
        [
            [(50, 130), (350, 130), (350, 160), (50, 160)],
            ["这是第二行内容，包含中文和英文", 0.96]
        ]
    ]
    
    # 原图尺寸（像素）- 实际使用时替换为你的图片尺寸
    img_width = 500
    img_height = 300
    
    # 生成Word文档
    ocr_results_to_word(
        ocr_results=ocr_results,
        output_path="ocr_result.docx",
        img_width=img_width,
        img_height=img_height
    )
