﻿
import datetime
import pandas as pd
from pathlib import Path

from docx import Document,opc,oxml
from docx.shared import Inches,Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
sys.path.append("..")
sys.path.append(".")

from __init__ import *
from uuid import uuid4
from PIL import Image
import os
import time
import json
from base.file_tools import read_write_async,read_write_sync,download_async,download_sync

from base.path_tools import normal_path
from __init__ import *
from base.com_log import logger ,record_detail,usage_time,logger_helper
from base import setting as setting
from base.string_tools import sanitize_filename
from docx.enum.style import WD_STYLE_TYPE
import asyncio
from DrissionPage.common import By
from DrissionPage import WebPage
from DrissionPage._elements.chromium_element import ChromiumElement

from lxml import etree
def convert_milliseconds_to_datetime(milliseconds):
    # 将毫秒时间戳转换为秒时间戳
    seconds = milliseconds / 1000.0
    
    # 将秒时间戳转换为 datetime 对象
    dt = datetime.datetime.fromtimestamp(seconds)
    
    # 格式化 datetime 对象为年月日时分秒格式
    formatted_date = dt.strftime('%Y-%m-%d %H:%M:%S')
    
    return formatted_date

def Num(thums):
    if type(thums)==str:
        thums=thums.strip()
    
        lst={"千":1000,"万":10000}
        scale:int=1
        
        for key in lst:
            if key in thums:
                scale*=lst[key]
                thums=thums.replace(key,"")
        
        return int(float(thums)*scale)  
    elif type(thums)==int:
        return thums
    else:
        return 0


def convert_image_to_jpg(image_path,dest_path=None):

    if not dest_path:
        dest_path=image_path

    
    helper=logger_helper("图片类型转换",f"{image_path}->{dest_path}")
    
    if not os.path.exists(image_path):
        helper.error("失败","图片文件不存在")
        return
    # 打开图片
    image = Image.open(image_path).convert('RGB')
    image.save(dest_path, 'JPEG')
    helper.trace("成功",update_time_type=True)

    

def add_hyperlink(paragraph, url, text, color=None):
    # 获取文档部分
    part = paragraph.part
    
    # 创建超链接关系
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # 创建超链接元素
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id)
    
    # 创建新的运行元素
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')
    
    # 设置颜色
    if color:
        c = oxml.shared.OxmlElement('w:color')
        c.set(oxml.shared.qn('w:val'), color)
        rPr.append(c)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph._p.add_r()
    r.append(hyperlink)

    return hyperlink
        
        
        


def wait_for_file_exist(file_path, timeout=5):
    start_time = time.time()
    cur_path = normal_path(file_path)
    target="等待文件存在"

    while time.time() - start_time < timeout:
        if os.path.exists(cur_path):
            logger.debug(record_detail(target,"成功",f"{usage_time(start_time)}"))
            return True
        time.sleep(0.5)  # 每0.5秒检查一次
    
    logger.error(record_detail(target,"失败",f"{usage_time(start_time)}"))
    return False

def url_file_suffix(url:str)->str:
    suffix=".xxyy"
    if url:
        parts=url.split("_")
        if len(parts)>2:
            suffix=parts[-2]
    return suffix
def is_jpg(file_path)->bool:
    item= Path(file_path)
    return item.suffix.lower() in [".jpg",".jpeg"]


class NoteInfo:
    #CountIndex=0
    def __init__(self,title,content,create_time,update_time,image_urls,author,thumbs,
                 collected,shared,comment,note_id,current_time,video_urls,image_lst=[],
                 video_lst=[],note_path="",type="note",link=""):
        self.title=title
        self.content=content
        self.create_time=create_time    
        self.update_time=update_time  
        self.image_urls=image_urls if isinstance(image_urls,list) else json.loads(image_urls)
        self.video_urls=video_urls if isinstance(video_urls,list) else json.loads(video_urls)
        
        self.author=author
        self.thumbs=thumbs
        self.collected=collected
        self.shared=shared
        self.comment=comment
        self.note_id=note_id
        self.current_time=current_time

        self.image_lst=image_lst if isinstance(image_lst,list) else json.loads(image_lst)
        self.video_lst=video_lst if isinstance(video_lst,list) else json.loads(video_lst)

        
        self.note_path=note_path
        self.type=type if type else ""
        self.link=link if link else ""
        
    def __str__(self) -> str:

        contents=[]
        if self.has_title:
            contents.append(f"title:{self.title}")
        if self.has_link:
            contents.append(f"link:{self.link}")
        if self.has_type:
            contents.append(f"type:{self.type}")
        
        if self.has_note_id:
            contents.append(f"note_id:{self.note_id}")
        if self.has_current_time:
            contents.append(f"current_time:{self.current_time}")
        if self.has_create_time:
            contents.append(f"create_time:{self.create_time}")
        if self.has_update_time:
            contents.append(f"update_time:{self.update_time}")
        if self.has_images:
            contents.append(f"image_urls:{"\n".join(self.image_urls)}")
        if self.has_video:
            contents.append(f"video_urls:{"\n".join(self.video_urls)}")
        if self.has_author:
            contents.append(f"author:{self.author}")
        if self.has_thumbs:
            contents.append(f"thumbs:{self.thumbs}")
        if self.has_collected:
            contents.append(f"collected:{self.collected}")
        if self.has_shared:
            contents.append(f"shared:{self.shared}")
        if self.has_comment:
            contents.append(f"comment:{self.comment}")
        if self.has_content:
            contents.append(f"content:{self.content}")
        
        if self.has_image_lst:
            contents.append(f"image_lst:{"\n".join(self.image_lst)}")
        if self.has_video_lst:
            contents.append(f"video_lst:{"\n".join(self.video_lst)}")
        
        return '\n'.join(contents)+"\n"


    @property
    def has_title(self)->bool:
        return self.title
    @property
    def has_images(self)->bool:
        return self.image_urls
    
    @property
    def has_video(self)->bool:
        return self.video_urls
    
    @property
    def has_content(self)->bool:
        return self.content
    
    @property
    def has_author(self)->bool:
        return self.author
    
    @property
    def has_thumbs(self)->bool:
        return self.thumbs
    
    @property
    def has_collected(self)->bool:
        return self.collected
    
    @property
    def has_shared(self)->bool:
        return self.shared
    
    @property
    def has_comment(self)->bool:
        return self.comment
    
    @property
    def has_note_id(self)->bool:
        return self.note_id
    
    @property
    def has_create_time(self)->bool:
        return self.create_time
    
    @property
    def has_update_time(self)->bool:
        return self.update_time
    
    @property
    def has_current_time(self)->bool:
        return self.current_time
    
    @property
    def has_image_lst(self)->bool:
        return self.image_lst
    @property
    def has_video_lst(self)->bool:
        return self.video_lst
    
    @property
    def has_type(self)->bool:
        return self.type
    @property
    def has_link(self)->bool:
        return self.link
    
    
    @property
    def DestImageLst(self)->list[str]:
        lst=[]
        for file_path in self.image_lst:
            if is_jpg(file_path):
                lst.append(file_path)
            else:
                cache_file=Path(file_path).with_suffix(".jpg")
                lst.append(str(cache_file))
        return lst
        
    @property
    def image_cache_dest_lst(self) ->list[tuple[str,str]]:
        lst=[]
        for cache,dest in zip(self.image_lst,self.DestImageLst):
            if cache!=dest:
                lst.append((cache,dest))
        return lst
    def set_root_dir(self,root_dir,cur_index:int =-1):


        sub_dir=f"{cur_index}_{self.title}" if cur_index>0 else self.title

        cur_dir=os.path.join(root_dir,sub_dir)

        dest_image_dir=os.path.join(cur_dir,"images/")
        dest_video_dir=os.path.join(cur_dir,"videos/")
        self.note_path=os.path.join(cur_dir,  f"{self.title}.txt")

        os.makedirs(os.path.dirname(dest_image_dir),exist_ok=True)
        os.makedirs(os.path.dirname(dest_video_dir),exist_ok=True)
        
        #原始的图片类型
        self.image_lst=[normal_path(os.path.join(dest_image_dir,f"{i+1}.{url_file_suffix(url)}")) for i,url in enumerate(self.image_urls)]
        self.video_lst=[normal_path(os.path.join(dest_video_dir,f"{i+1}.mp4")) for i in range(len(self.video_urls))] 
        
    #下载图片及视频
    async def download(self):
       
        image_cache_dest_lst=self.image_cache_dest_lst
        convert_cache=[cache for cache,_ in image_cache_dest_lst]
        
        async def cur_download_async(self,url,dest):
            await download_async(url,dest)
            
            #若是 图片类型不是 jpg，则转换为jpg
            if dest in convert_cache:
                convert_image_to_jpg(dest,image_cache_dest_lst[convert_cache.index(dest)][1])
            

        #图片 + 视频
        tasks=[ cur_download_async(self,url,dest) for url,dest in zip(self.image_urls+self.video_urls,self.image_lst+self.video_lst) ]
        return tasks
        # await asyncio.gather(*tasks)
          
    #文本
    async def write_to_notepad(self):
        await read_write_async(str(self),self.note_path,mode="w",encoding="utf-8")
        logger.debug(record_detail("写入记事本","成功", f"{self.note_path}"))
    
    #异步下载以及写入到文本
    async def handle_note(self):
        image_cache_dest_lst=self.image_cache_dest_lst
        convert_cache=[cache for cache,_ in image_cache_dest_lst]
        
        async def cur_download_async(self,url,dest):
            await download_async(url,dest)
            
            #若是 图片类型不是 jpg，则转换为jpg
            if dest in convert_cache:
                convert_image_to_jpg(dest,image_cache_dest_lst[convert_cache.index(dest)][1])
            

        #图片 + 视频
        tasks=[ cur_download_async(self,url,dest) for url,dest in zip(self.image_urls+self.video_urls,self.image_lst+self.video_lst) ]
        tasks.append(self.write_to_notepad())
        await asyncio.gather(*tasks)
    
    
    def write_to_word(self,document:Document):

        start_time=time.time()
        target=f"添加文档word标题:{self.title}"
        
        logger.trace(record_detail(target,"开始",""))
        try:
            heading = document.add_heading(self.title, 2)
            heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加表格
            
            def merge_val(ls):
                return ":".join(list(map(str,ls)) )
            
            def merge_vals(ls):
                return " ".join(list(map(merge_val,ls)) )
            
            thumbs=[
            ["点赞",self.thumbs],
            ["收藏",self.collected],
            ["分享",self.shared],
            ]
            
            
            ls=[

                [merge_val(["create",self.create_time]),merge_val(["作者",self.author]),],
                [merge_val(["update",self.update_time]),merge_vals(thumbs),],
                [merge_val(["now",self.current_time]),merge_val(["类型",self.type]),]
            ]
            
            row_count=len(ls)
            #表格
            table = document.add_table(rows=row_count, cols=2)

            for row in range(row_count):
                for cell in range(2):
                    table.rows[row].cells[cell].text=ls[row][cell]

            #添加note_id,附带超链接
            id_paragraph = document.add_paragraph()
            id_info=f"id:{self.note_id}"
            #添加链接
            if self.has_link:
                add_hyperlink(id_paragraph, self.link,id_info, color='0563C1')
            else:
                id_paragraph.add_run(id_info)
            
            #图片
            if self.DestImageLst:
                pic_paragraph = document.add_paragraph()
                pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for image_path in self.DestImageLst:
                    cur_path=normal_path(image_path)
                    if not (os.path.exists(cur_path) or wait_for_file_exist(cur_path,5)):
                        logger.error( record_detail("获取文件","失败", detail=f"image_path:{cur_path} 文件不存在"))
                        continue
                    
                    try:
                        pic_paragraph.add_run().add_picture( cur_path,width=Cm(7))
                    except:
                        logger.error(record_detail("插入图片","失败",f"image_path:{cur_path} 插入word文档【{self.title}】失败"))
                        continue

            
            # if self.video_lst:
            #     video_paragraph = document.add_paragraph().add_run()
            #     for video_path in self.video_lst:
            #         video_paragraph.add_video(video_path,width=Inches(1.5))
            #         video_paragraph.add_run()
                
            
            if self.has_content:
                document.add_paragraph(self.content)
        except Exception as e:
            logger.error(record_detail(target,"失败",f"异常:{e},{usage_time(start_time)}"))
            # raise e
            
        document.add_paragraph('')
        document.add_paragraph('')
        
        
        
        logger.debug(record_detail(target,"成功", usage_time(start_time)))
        pass


class Section:
    def __init__(self,sec,yVal,id,already,error_count=0):
        # self.sec:ChromiumElement=sec
        self.sec=sec
        self.yVal=yVal
        self.id=id
        self.already=already
        self.error_count=error_count
    def id(self):
        return self.id
    
    @property
    def url(self):

        root = etree.HTML(self.sec.inner_html)
        # 使用 XPath 提取 href 属性值
        href_value = root.xpath('//a[@class="cover ld mask"]/@href')
        return f"https://www.xiaohongshu.com{href_value[0]}" if href_value else ""

def contain_search_key(str):
    lst=["大家都在搜","相关搜索","热门搜索"]
    return  any([key in str for key in lst])
    

class SectionManager:
    def __init__(self,wp:WebPage):
        self.wp=wp

        self.secs=[]
        self.cur_id=""

        
    def __set_secs(self,secs:list[Section]):
        self.secs=secs
        
        # for sec in self.cur_secs:
        #     sec_ids=[sec.id for sec in self.secs]
        #     if sec.id not in sec_ids:
        #         self.secs.append(sec)
        #     else:
        #         index=sec_ids.index(sec.id)
        #         self.secs[index]=sec
        
    @property
    def urls(self):
        lst= [sec.url for sec in self.secs if sec]
        return list(filter(lambda x:x,lst))

    @property
    def last(self):
        return self.secs[-1].sec if self.secs else None
    
    def update(self):
        target=(By.XPATH,'//section[@class="note-item"]')
        if not self.wp.wait.eles_loaded(target):
            return
            
        org_secs=self.wp.eles(target)
        sorted(org_secs,key=lambda x:x.rect.midpoint[1])

        
        secs=[Section(sec,sec.rect.midpoint[1],sec.raw_text,False,0)   for sec in org_secs if not contain_search_key(sec.raw_text) ]
        

 
        
        if self.secs:
            for sec in secs:
                org=self.get_by_id(sec.id)
                if org:
                    sec.already=org.already
                    sec.error_count=org.error_count
                    

        self.__set_secs(secs)
    def get_by_id(self ,id):
        secs=[sec for sec in self.secs if sec.id==id]
        return secs[0] if secs else None
        
    @property
    def ids(self):
        return [sec.id for sec in self.secs]
    
    

    def next(self):

        for sec in self.secs:
            if not sec.already:
                sec.already=True
                self.cur_id=sec.id
                yield sec.sec
        yield None

     
    def resume_cur(self):
        sec=self.get_by_id(self.cur_id)
        if sec:
            sec.error_count+=1
            if sec.error_count>3:
                print(f"{self.cur_id} 失败次数已超3次")
                sec.already=True
                return
            sec.already=False


        
        
    def count(self):

        return sum([ 0 if sec.already else 1  for sec in self.secs ])
    
def to_theme_word(theme_name,root_dir,dict_data):
        # 创建一个新的文档
    document = Document()
    start_time=time.time()
    word_path=os.path.join(root_dir,f"{theme_name}.docx")
    target=f"写入文档{word_path}"
    
    logger.debug(record_detail(target,"开始", "..."))
    try:
        #整理到word中
        for note_info in dict_data:
            info=NoteInfo(**note_info)
            info.write_to_word(document)
        document.save(word_path)
    except Exception as e:
        logger.error(record_detail(target,"失败", detail=f"{str(e)},{usage_time(start_time)}"))
        # raise e
        return
        
        
        
    logger.debug(record_detail(target,"成功", usage_time(start_time)))
    

if __name__ == '__main__':
    convert_image_to_jpg(r"F:\手工作业\一家人\放风筝.jfif",r"F:\手工作业\一家人\放风筝.jpg")