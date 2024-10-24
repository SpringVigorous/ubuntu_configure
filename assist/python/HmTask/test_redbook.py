from queue import Queue
from base_task import ThreadTask
import threading

from DrissionPage import WebPage
from DrissionPage.common import Actions
from pprint import pprint
import os
import time
import json
from requests.structures import CaseInsensitiveDict
import requests
import sys


from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC


from __init__ import *
from base.com_log import logger as logger
from base import setting as setting
from base.string_tools import sanitize_filename
# from base.cookies_tools import save_cookies,load_cookies,exist_cookies
from base.file_tools import read_write_async,read_write_sync,download_async

from base.path_tools import normal_path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from HmTask.redbook_tools import *




        

class Interact():
    def __init__(self, theme_lst:list,root_dir:str=setting.redbook_notes_dir,search_count:int=200) -> None:
        self.theme_lst=theme_lst
        self.office_event=threading.Event()

        self.jsons=[]
        self.wp=WebPage()
        self.root_dir=root_dir
        self.search_count=search_count
        url='https://www.xiaohongshu.com/'
        
        self.wp.get(url)
        

        time.sleep(5) #初次等待登录
        
    def run(self):
        for theme  in self.theme_lst:

            if theme is None:
                continue
            output_data = self.handle_data(theme)
            if not output_data:
                continue
            
            
            data_queue=Queue()
            data_queue.put(output_data)
            stop_event=threading.Event()
            write_excel=WriteExcel(input_queue=data_queue,stop_event=stop_event)
            write_word=WriteWord(input_queue=data_queue,stop_event=stop_event)
            write_excel.start()
            write_word.start()
            data_queue.join()
            stop_event.set()
            write_excel.join()
            write_word.join()
        
    def handle_data(self, data):
        
        search_input = self.wp.ele('xpath://input[@class="search-input"]')
        search_input.clear()
        search_input.input(f'{data}\n')
        time.sleep(.4)

        
        seach_button=self.wp.ele('xpath://div[@class="search-icon"]')
        if not seach_button:
            sys.exit(0)
        self.wp.ele('xpath://div[@class="search-icon"]').click()
        self.wp.listen.start(['web/v1/search/notes',"api/sns/web/v1/feed"]) 
        packet = self.wp.listen.wait()



        secManager=SectionManager(self.wp)
        secManager.update()
        
        
        json_queue=Queue()
        data_queue=Queue()
        datas_queue=Queue()
        
        stop_event=threading.Event()
        parse= Parse(json_queue,output_queue=data_queue,datas_queue=datas_queue,stop_envent=stop_event)
        parse.start()
        
        datas=[]
        while (secManager.count())>0:
            
            if len(datas)>self.search_count:
                self.stoped=True
                break
            sec=next(secManager.next()) 
            if not sec:
                continue
            time.sleep(.1)
            
            try:
                sec.click()
            except:
                secManager.resume_cur()
                secManager.update()
                continue
            pack=self.wp.listen.wait()
            body=pack.response.body
            if body:
                json_queue.put(body)    #整理到内部队列
                datas.append(body)

            close_flag=self.wp.ele('xpath://div[@class="close close-mask-dark"]')
            if close_flag:
                try:
                    close_flag.click()
                except:
                    continue
                
        datas_queue.put(datas)
                
        json_queue.join()
        Parse.set()
        parse.join()
        return datas

def convert_milliseconds_to_datetime(milliseconds):
    # 将毫秒时间戳转换为秒时间戳
    seconds = milliseconds / 1000.0
    
    # 将秒时间戳转换为 datetime 对象
    dt = datetime.datetime.fromtimestamp(seconds)
    
    # 格式化 datetime 对象为年月日时分秒格式
    formatted_date = dt.strftime('%Y-%m-%d %H:%M:%S')
    
    return formatted_date

def Num(thums):
    if type(thums)==str:
        thums=thums.strip()
    
        lst={"千":1000,"万":10000}
        scale:int=1
        
        for key in lst:
            if key in thums:
                scale*=lst[key]
                thums=thums.replace(key,"")
        
        return int(float(thums)*scale)  
    elif type(thums)==int:
        return thums
    else:
        return 0

from PIL import Image
def convert_image_to_jpg(image_path,dest_path=None):
    if not dest_path:
        dest_path=image_path
    if not os.path.exists(image_path):
        logger.error(f"图片文件不存在:{image_path},{image_path}->{dest_path}失败")
        return
    # 打开图片
    image = Image.open(image_path).convert('RGB')
    # 保存为 JPEG 格式


    image.save(dest_path, 'JPEG')
    if os.path.exists(dest_path):
        os.remove(image_path)
    
        return
    # 打开图片
    image = Image.open(image_path).convert('RGB')
    # 保存为 JPEG 格式
    if not dest_path:
        dest_path=image_path

    image.save(dest_path, 'JPEG')
    

class NoteInfo:
    #CountIndex=0
    def __init__(self,title,content,create_time,update_time,image_urls,author,thumbs,collected,shared,comment,note_id,current_time,video_urls,image_lst=[],video_lst=[],note_path=""):
        self.title=title
        self.content=content
        self.create_time=create_time    
        self.update_time=update_time  
        self.image_urls=image_urls if isinstance(image_urls,list) else json.loads(image_urls)
        self.video_urls=video_urls if isinstance(video_urls,list) else json.loads(video_urls)
        
        self.author=author
        self.thumbs=thumbs
        self.collected=collected
        self.shared=shared
        self.comment=comment
        self.note_id=note_id
        self.current_time=current_time

        self.image_lst=image_lst if isinstance(image_lst,list) else json.loads(image_lst)
        self.video_lst=video_lst if isinstance(video_lst,list) else json.loads(video_lst)

        self.note_path=note_path
        
        
    def __str__(self) -> str:

        contents=[]
        if self.has_title:
            contents.append(f"title:{self.title}")
        if self.has_note_id:
            contents.append(f"note_id:{self.note_id}")
        if self.has_current_time:
            contents.append(f"current_time:{self.current_time}")
        if self.has_create_time:
            contents.append(f"create_time:{self.create_time}")
        if self.has_update_time:
            contents.append(f"update_time:{self.update_time}")
        if self.has_images:
            contents.append(f"image_urls:{"\n".join(self.image_urls)}")
        if self.has_video:
            contents.append(f"video_urls:{"\n".join(self.video_urls)}")
        if self.has_author:
            contents.append(f"author:{self.author}")
        if self.has_thumbs:
            contents.append(f"thumbs:{self.thumbs}")
        if self.has_collected:
            contents.append(f"collected:{self.collected}")
        if self.has_shared:
            contents.append(f"shared:{self.shared}")
        if self.has_comment:
            contents.append(f"comment:{self.comment}")
        if self.has_content:
            contents.append(f"content:{self.content}")
        
        if self.has_image_lst:
            contents.append(f"image_lst:{"\n".join(self.image_lst)}")
        if self.has_video_lst:
            contents.append(f"video_lst:{"\n".join(self.video_lst)}")
        
        return '\n'.join(contents)+"\n"

    @property
    def has_title(self)->bool:
        return self.title
    @property
    def has_images(self)->bool:
        return self.image_urls
    
    @property
    def has_video(self)->bool:
        return self.video_urls
    
    @property
    def has_content(self)->bool:
        return self.content
    
    @property
    def has_author(self)->bool:
        return self.author
    
    @property
    def has_thumbs(self)->bool:
        return self.thumbs
    
    @property
    def has_collected(self)->bool:
        return self.collected
    
    @property
    def has_shared(self)->bool:
        return self.shared
    
    @property
    def has_comment(self)->bool:
        return self.comment
    
    @property
    def has_note_id(self)->bool:
        return self.note_id
    
    @property
    def has_create_time(self)->bool:
        return self.create_time
    
    @property
    def has_update_time(self)->bool:
        return self.update_time
    
    @property
    def has_current_time(self)->bool:
        return self.current_time
    
    @property
    def has_image_lst(self)->bool:
        return self.image_lst
    @property
    def has_video_lst(self)->bool:
        return self.video_lst
    
    def set_root_dir(self,root_dir,cur_index:int =-1):
        # index=cur_index
        # if cur_index<1:
        #     NoteInfo.CountIndex+=1
        #     index= NoteInfo.CountIndex
            

        # sub_dir=f"{index}"

        sub_dir=f"{cur_index}_{self.title}" if cur_index>0 else self.title

        cur_dir=os.path.join(root_dir,sub_dir)

        dest_image_dir=os.path.join(cur_dir,"images/")
        dest_video_dir=os.path.join(cur_dir,"videos/")
        self.note_path=os.path.join(cur_dir,  f"{self.title}.txt")

        os.makedirs(os.path.dirname(dest_image_dir),exist_ok=True)
        os.makedirs(os.path.dirname(dest_video_dir),exist_ok=True)

        self.image_lst=[normal_path(os.path.join(dest_image_dir,f"{i+1}.jpg")) for i in range(len(self.image_urls))]
        self.video_lst=[normal_path(os.path.join(dest_video_dir,f"{i+1}.mp4")) for i in range(len(self.video_urls))] 
        
          
    async def write_to_notepad(self):

        urls=self.image_urls+self.video_urls
        
        #查找前缀
        suffix=".xxyy"
        if self.image_urls:
            url=self.image_urls[0]
            parts=url.split("_")
            if len(parts)>2:
                suffix=parts[-2]
        
        image_lst=[Path(file_path).with_suffix("."+suffix)  for file_path in self.image_lst]
        
        dests=image_lst+self.video_lst
        #图片 + 视频
        tasks=[download_async(url,dest) for url,dest in zip(urls,dests) ]
        #文本
        tasks.append(read_write_async(str(self),self.note_path,mode="w",encoding="utf-8"))
        
        await asyncio.gather(*tasks)
        
        
        
        if suffix in ["jpg","jpeg"]:
            return
        for file_path,dest_path in zip(image_lst,self.image_lst):
            convert_image_to_jpg(file_path,dest_path)
        

    def to_word(self,document:Document):

        document.add_heading(self.title, 2)
        # 添加表格
        ls=[
        ["id",self.note_id],
        ["作者",self.author],
        ["点赞",self.thumbs],
        ["当前时间",self.current_time],
        ["收藏",self.collected],
        ["分享",self.shared],

        ["创建时间",self.create_time],
        ["更新时间",self.update_time],
        ]

        def TableVal(index:int):
            if index<len(ls):
                return f"{ls[index][0]}:{ls[index][1]}"
            else:
                return ""
        
        #表格
        table = document.add_table(rows=3, cols=3)
        k=0
        for row in range(3):
            for cell in range(3):
                table.rows[row].cells[cell].text=TableVal(k)
                k+=1
                if k>len(ls):
                    break
        
        
        
        #图片
        if self.image_lst:
            pic_paragraph = document.add_paragraph()
            for image_path in self.image_lst:
                cur_path=normal_path(image_path)
                try:
                    pic_paragraph.add_run().add_picture( cur_path,width=Cm(7))
                except:
                    logger.error(f"image_path:{cur_path} 插入word文档【{self.title}】失败")
                    continue

        
        # if self.video_lst:
        #     video_paragraph = document.add_paragraph().add_run()
        #     for video_path in self.video_lst:
        #         video_paragraph.add_video(video_path,width=Inches(1.5))
        #         video_paragraph.add_run()
            
        
        
        document.add_paragraph(self.content)
        pass



class Parse(ThreadTask):
    def __init__(self, input_queue, output_queue,stop_event, datas_queue) -> None:
        super().__init__(input_queue=input_queue, output_queue=output_queue,stop_event=stop_event)
        self.datas_queue=datas_queue
        self.datas=[]
    def _handle_data(self, raw_data):
        if not "data" in raw_data:
            return None
        
        note_data=raw_data["data"]
        note_info=note_data["items"][0]
        
        id=note_info["id"]
        model_type=note_info["model_type"]
        note_card=note_info["note_card"]
        topics=[ tag["name"] for tag in note_card["tag_list"]] if "tag_list" in note_card else []
        if not topics:
            return None
        
        content=note_card["desc"] if "desc" in note_card else ""
        user=note_card["user"]
        user_id=user["user_id"]
        user_name=user["nickname"]

        user_icon=user["avatar_url"] if "avatar_url" in user else ""
        interact_info=note_card["interact_info"]
        liked_count=Num(interact_info.get("liked_count"))
        collected_count=Num(interact_info.get("collected_count",0))
        share_count=Num(interact_info.get("share_count",0))
        comment_count=note_info.get("comment_count",0)
        
        image_urls=[item["url_default"]  for item in note_card["image_list"]]  if "image_list" in note_card else []
        create_time=convert_milliseconds_to_datetime(note_card["time"])
        update_time=convert_milliseconds_to_datetime(note_card["last_update_time"])
        note_id=note_card["note_id"]
        title=sanitize_filename(note_card["title"])
        if not title:
            title=str(uuid4())
        current_time=convert_milliseconds_to_datetime(note_data["current_time"])
        video_lst=[video["master_url"] for video in note_card["video"]["media"]["stream"]["h264"]] if "video" in note_card else []

        return  NoteInfo(title=title,
                        content=content,
                        create_time=create_time,
                        update_time=update_time,
                        image_urls=image_urls,
                        author=user_name,
                        thumbs=liked_count,
                        collected=collected_count,
                        shared=share_count,
                        comment=comment_count,
                        note_id=note_id,
                        current_time=current_time,
                        video_urls=video_lst)

class InputTask(ThreadTask):
    def __init__(self, input_queue, stop_event) -> None:
            super().__init__(input_queue=input_queue, stop_event=stop_event)
    def _handle_data(self, data):
        pass

class DownLoad(InputTask):
    def _handle_data(self, data):
        pass

class WriteNotePad(InputTask):

    def _handle_data(self, data):
        pass

#以下两个可以并行
class WriteExcel(InputTask):

    def _handle_data(self, data):
        pass

class WriteWord(InputTask):
    def _handle_data(self, data):
        pass




# class App:
#     def __init__(self,themes:list) -> None:
        
        
#         self.theme_queue=Queue()
        
        

#         self.json_queue=Queue()
#         self.data_queue=Queue()
#         self.datas_queue=Queue()       
        
#         self.parse_event=threading.Event()
#         self.download_event=threading.Event()
#         self.notepad_envent=threading.Event()

#         self.interact=Interact(input_queue=self.theme_queue,output_queue=self.json_queue)
#         self.parse=Parse(input_queue=self.json_queue,output_queue=self.data_queue,stop_event=self.parse_event,datas_queue=self.datas_queue) 
#         self.download=DownLoad(input_queue=self.datas_queue,stop_event=self.download_event)
#         self.write_notepad=WriteNotePad(input_queue=self.data_queue,stop_event=self.download_event)
        
        

        
        
    
#     def run(self):
#         for theme in self.themes:
#             self.theme_queue.put(theme)
#         self.interact.start()
#         self.theme_queue().join()
#         self.parse_event.set()
        
    
#     pass



# if __name__ == '__main__':
#     app=App()
#     app.run()