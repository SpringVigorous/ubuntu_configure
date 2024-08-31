from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_docx():
    # 创建一个新的文档
    document = Document()

    # 添加标题
    document.add_heading('示例文档', 0)

    # 添加段落
    p = document.add_paragraph('这是一个示例段落。')
    p.add_run('这是段落的一部分。').bold = True
    p.add_run('这是段落的另一部分。').italic = True

    # 添加另一个段落
    document.add_paragraph('这是另一个段落。')

    # 添加有序列表
    list_paragraph = document.add_paragraph()
    list_paragraph.add_run('项目 1').bold = True
    list_paragraph.add_run(' - 这是项目 1 的描述。')
    list_paragraph = document.add_paragraph()
    list_paragraph.add_run('项目 2').bold = True
    list_paragraph.add_run(' - 这是项目 2 的描述。')
    
    

    # 添加图片
    pic_paragraph = document.add_paragraph()
    pic=pic_paragraph.add_run().add_picture(r'D:\Document\Tencent\Wechat\WeChat Files\wxid_opokem0kmeqf22\FileStorage\Cache\2024-08\35c7d5e055318963c8b79cb824d0a4e9_t.jpg', width=Inches(1.25))
    pic_paragraph.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    # pic.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    
    latter=document.add_paragraph("Hello,")
    latter.add_run("Kitty").bold=True
    latter.add_run("!").italic=True
    latter.paragraph_format.space_after=Inches(8)
    latter.paragraph_format.space_before=Inches(10)
    # latter.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 添加表格
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '列 1'
    hdr_cells[1].text = '列 2'
    hdr_cells[2].text = '列 3'

    for i in range(3):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = '数据 ' + str(i)
        row_cells[2].text = '更多数据 ' + str(i)

    # 保存文档
    document.save('example.docx')

def read_docx():
    doc=Document('example1.docx')
    for i,paragraph in enumerate(doc.paragraphs):
        print(i+1,paragraph.text,paragraph.style.name)
    
    
if __name__ == '__main__':
    create_docx()
    # read_docx()