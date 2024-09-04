from DrissionPage import WebPage
from DrissionPage.common import Actions
from pprint import pprint
import os
import time
import json
from requests.structures import CaseInsensitiveDict
import requests
import sys


from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC


from __init__ import *
from base.com_log import logger as logger
from base import setting as setting
from base.string_tools import sanitize_filename
# from base.cookies_tools import save_cookies,load_cookies,exist_cookies
from base.file_tools import read_write_async,read_write_sync,download_async
from base.com_log import logger as logger
from base.path_tools import normal_path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


            
import asyncio 
import aiofiles
import aiohttp
        
        
import datetime
import pandas as pd
from pathlib import Path

from docx import Document
from docx.shared import Inches,Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from uuid import uuid4


def convert_milliseconds_to_datetime(milliseconds):
    # 将毫秒时间戳转换为秒时间戳
    seconds = milliseconds / 1000.0
    
    # 将秒时间戳转换为 datetime 对象
    dt = datetime.datetime.fromtimestamp(seconds)
    
    # 格式化 datetime 对象为年月日时分秒格式
    formatted_date = dt.strftime('%Y-%m-%d %H:%M:%S')
    
    return formatted_date

def Num(thums):
    if type(thums)==str:
        thums=thums.strip()
    
        lst={"千":1000,"万":10000}
        scale:int=1
        
        for key in lst:
            if key in thums:
                scale*=lst[key]
                thums=thums.replace(key,"")
        
        return int(float(thums)*scale)  
    elif type(thums)==int:
        return thums
    else:
        return 0

from PIL import Image
def convert_image_to_jpg(image_path,dest_path=None):
    if not dest_path:
        dest_path=image_path
    if not os.path.exists(image_path):
        logger.error(f"图片文件不存在:{image_path},{image_path}->{dest_path}失败")
        return
    # 打开图片
    image = Image.open(image_path).convert('RGB')
    # 保存为 JPEG 格式


    image.save(dest_path, 'JPEG')
    if os.path.exists(dest_path):
        os.remove(image_path)
    
        return
    # 打开图片
    image = Image.open(image_path).convert('RGB')
    # 保存为 JPEG 格式
    if not dest_path:
        dest_path=image_path

    image.save(dest_path, 'JPEG')
    

class NoteInfo:
    #CountIndex=0
    def __init__(self,title,content,create_time,update_time,image_urls,author,thumbs,collected,shared,comment,note_id,current_time,video_urls,image_lst=[],video_lst=[],note_path=""):
        self.title=title
        self.content=content
        self.create_time=create_time    
        self.update_time=update_time  
        self.image_urls=image_urls if isinstance(image_urls,list) else json.loads(image_urls)
        self.video_urls=video_urls if isinstance(video_urls,list) else json.loads(video_urls)
        
        self.author=author
        self.thumbs=thumbs
        self.collected=collected
        self.shared=shared
        self.comment=comment
        self.note_id=note_id
        self.current_time=current_time

        self.image_lst=image_lst if isinstance(image_lst,list) else json.loads(image_lst)
        self.video_lst=video_lst if isinstance(video_lst,list) else json.loads(video_lst)

        self.note_path=note_path
        
        
    def __str__(self) -> str:

        contents=[]
        if self.has_title:
            contents.append(f"title:{self.title}")
        if self.has_note_id:
            contents.append(f"note_id:{self.note_id}")
        if self.has_current_time:
            contents.append(f"current_time:{self.current_time}")
        if self.has_create_time:
            contents.append(f"create_time:{self.create_time}")
        if self.has_update_time:
            contents.append(f"update_time:{self.update_time}")
        if self.has_images:
            contents.append(f"image_urls:{"\n".join(self.image_urls)}")
        if self.has_video:
            contents.append(f"video_urls:{"\n".join(self.video_urls)}")
        if self.has_author:
            contents.append(f"author:{self.author}")
        if self.has_thumbs:
            contents.append(f"thumbs:{self.thumbs}")
        if self.has_collected:
            contents.append(f"collected:{self.collected}")
        if self.has_shared:
            contents.append(f"shared:{self.shared}")
        if self.has_comment:
            contents.append(f"comment:{self.comment}")
        if self.has_content:
            contents.append(f"content:{self.content}")
        
        if self.has_image_lst:
            contents.append(f"image_lst:{"\n".join(self.image_lst)}")
        if self.has_video_lst:
            contents.append(f"video_lst:{"\n".join(self.video_lst)}")
        
        return '\n'.join(contents)+"\n"

    @property
    def has_title(self)->bool:
        return self.title
    @property
    def has_images(self)->bool:
        return self.image_urls
    
    @property
    def has_video(self)->bool:
        return self.video_urls
    
    @property
    def has_content(self)->bool:
        return self.content
    
    @property
    def has_author(self)->bool:
        return self.author
    
    @property
    def has_thumbs(self)->bool:
        return self.thumbs
    
    @property
    def has_collected(self)->bool:
        return self.collected
    
    @property
    def has_shared(self)->bool:
        return self.shared
    
    @property
    def has_comment(self)->bool:
        return self.comment
    
    @property
    def has_note_id(self)->bool:
        return self.note_id
    
    @property
    def has_create_time(self)->bool:
        return self.create_time
    
    @property
    def has_update_time(self)->bool:
        return self.update_time
    
    @property
    def has_current_time(self)->bool:
        return self.current_time
    
    @property
    def has_image_lst(self)->bool:
        return self.image_lst
    @property
    def has_video_lst(self)->bool:
        return self.video_lst
    
    def set_root_dir(self,root_dir,cur_index:int =-1):
        # index=cur_index
        # if cur_index<1:
        #     NoteInfo.CountIndex+=1
        #     index= NoteInfo.CountIndex
            

        # sub_dir=f"{index}"

        sub_dir=f"{cur_index}_{self.title}" if cur_index>0 else self.title

        cur_dir=os.path.join(root_dir,sub_dir)

        dest_image_dir=os.path.join(cur_dir,"images/")
        dest_video_dir=os.path.join(cur_dir,"videos/")
        self.note_path=os.path.join(cur_dir,  f"{self.title}.txt")

        os.makedirs(os.path.dirname(dest_image_dir),exist_ok=True)
        os.makedirs(os.path.dirname(dest_video_dir),exist_ok=True)

        self.image_lst=[normal_path(os.path.join(dest_image_dir,f"{i+1}.jpg")) for i in range(len(self.image_urls))]
        self.video_lst=[normal_path(os.path.join(dest_video_dir,f"{i+1}.mp4")) for i in range(len(self.video_urls))] 
        
          
    async def write_to_notepad(self):

        urls=self.image_urls+self.video_urls
        
        #查找前缀
        suffix=".xxyy"
        if self.image_urls:
            url=self.image_urls[0]
            parts=url.split("_")
            if len(parts)>2:
                suffix=parts[-2]
        
        image_lst=[Path(file_path).with_suffix("."+suffix)  for file_path in self.image_lst]
        
        dests=image_lst+self.video_lst
        #图片 + 视频
        tasks=[download_async(url,dest) for url,dest in zip(urls,dests) ]
        #文本
        tasks.append(read_write_async(str(self),self.note_path,mode="w",encoding="utf-8"))
        
        await asyncio.gather(*tasks)
        
        
        
        if suffix in ["jpg","jpeg"]:
            return
        for file_path,dest_path in zip(image_lst,self.image_lst):
            convert_image_to_jpg(file_path,dest_path)
        

    def to_word(self,document:Document):

        document.add_heading(self.title, 2)
        # 添加表格
        ls=[
        ["id",self.note_id],
        ["作者",self.author],
        ["点赞",self.thumbs],
        ["当前时间",self.current_time],
        ["收藏",self.collected],
        ["分享",self.shared],

        ["创建时间",self.create_time],
        ["更新时间",self.update_time],
        ]

        def TableVal(index:int):
            if index<len(ls):
                return f"{ls[index][0]}:{ls[index][1]}"
            else:
                return ""
        
        #表格
        table = document.add_table(rows=3, cols=3)
        k=0
        for row in range(3):
            for cell in range(3):
                table.rows[row].cells[cell].text=TableVal(k)
                k+=1
                if k>len(ls):
                    break
        
        
        
        #图片
        if self.image_lst:
            pic_paragraph = document.add_paragraph()
            for image_path in self.image_lst:
                cur_path=normal_path(image_path)
                try:
                    pic_paragraph.add_run().add_picture( cur_path,width=Cm(7))
                except:
                    logger.error(f"image_path:{cur_path} 插入word文档【{self.title}】失败")
                    continue

        
        # if self.video_lst:
        #     video_paragraph = document.add_paragraph().add_run()
        #     for video_path in self.video_lst:
        #         video_paragraph.add_video(video_path,width=Inches(1.5))
        #         video_paragraph.add_run()
            
        
        
        document.add_paragraph(self.content)
        pass

async def ParseOrgNote(raw_data):
    # raw_data=json.loads(body)
    if not "data" in raw_data:
        return None
    
    await asyncio.sleep(.1)
    
    note_data=raw_data["data"]
    note_info=note_data["items"][0]
    
    id=note_info["id"]
    model_type=note_info["model_type"]
    note_card=note_info["note_card"]
    topics=[ tag["name"] for tag in note_card["tag_list"]] if "tag_list" in note_card else []
    if not topics:
        return None
    
    content=note_card["desc"] if "desc" in note_card else ""
    user=note_card["user"]
    user_id=user["user_id"]
    user_name=user["nickname"]

    user_icon=user["avatar_url"] if "avatar_url" in user else ""
    interact_info=note_card["interact_info"]
    liked_count=Num(interact_info.get("liked_count"))
    collected_count=Num(interact_info.get("collected_count",0))
    share_count=Num(interact_info.get("share_count",0))
    comment_count=note_info.get("comment_count",0)
    
    image_urls=[item["url_default"]  for item in note_card["image_list"]]  if "image_list" in note_card else []
    create_time=convert_milliseconds_to_datetime(note_card["time"])
    update_time=convert_milliseconds_to_datetime(note_card["last_update_time"])
    note_id=note_card["note_id"]
    title=sanitize_filename(note_card["title"])
    if not title:
        title=str(uuid4())
    current_time=convert_milliseconds_to_datetime(note_data["current_time"])
    video_lst=[video["master_url"] for video in note_card["video"]["media"]["stream"]["h264"]] if "video" in note_card else []

    return  NoteInfo(title=title,
                    content=content,
                    create_time=create_time,
                    update_time=update_time,
                    image_urls=image_urls,
                    author=user_name,
                    thumbs=liked_count,
                    collected=collected_count,
                    shared=share_count,
                    comment=comment_count,
                    note_id=note_id,
                    current_time=current_time,
                    video_urls=video_lst)


class Section:
    def __init__(self,sec,yVal,id,already):
        self.sec=sec
        self.yVal=yVal
        self.id=id
        self.already=already

def contain_search_key(str):
    lst=["大家都在搜","相关搜索","热门搜索"]
    return  any([key in str for key in lst])
    

class SectionManager:
    def __init__(self,wp):
        self.wp=wp
        self.cur_secs=[]
        self.secs=[]
        self.cur_index=0
        
    def __set_secs(self,secs:list):
        self.cur_secs=secs
        
        for sec in self.cur_secs:
            sec_ids=[sec.id for sec in self.secs]
            if sec.id not in sec_ids:
                self.secs.append(sec)
            else:
                self.secs[sec_ids.index(sec.id)]=sec
        


    def update(self):
        
        # await asyncio.sleep(.3)
        time.sleep(.3)
        
        org_secs=self.wp.eles("xpath://section")
        sorted(org_secs,key=lambda x:x.rect.midpoint[1])

        
        secs=[Section(sec,sec.rect.midpoint[1],sec.raw_text,False)   for sec in org_secs if not contain_search_key(sec.raw_text) ]
        
 
        
        if self.secs:
            for sec in secs:
                
                org=self.get_by_id(sec.id)
                if org:
                    sec.already=org.already
                    
        
        

        self.__set_secs(secs)
    def get_by_id(self ,id):
        secs=[sec for sec in self.secs if sec.id==id]
        return secs[0] if secs else None
        

    
    def next(self):

        for sec in self.cur_secs:
            if not sec.already:
                sec.already=True
                self.cur_index=self.cur_secs.index(sec)
                yield sec.sec
     
    def resume_cur(self):
        if self.cur_index<len(self.cur_secs):
            self.cur_secs[self.cur_index].already=False
        
        
    def count(self):

        return sum([ 0 if sec.already else 1  for sec in self.cur_secs ])


        


class RedBookSearch:
    is_first=True
    def __init__(self,wp:WebPage,theme_name:str,root_dir:str=setting.redbook_notes_dir,search_count:int=200,queue_count:int=100) -> None:
        self.wp=wp
        self.theme_name=theme_name
        self.root_dir=root_dir
        self.search_count=search_count
        self.data_queue=asyncio.Queue(maxsize=queue_count)
        self.cur_index=0
        self.stoped=False
        
        
        pass
    @property
    def ThemeDir(self):
        return os.path.join(self.root_dir,self.theme_name)
    @property
    def ThemeHistoryDir(self):
        cur_dir= os.path.join(self.ThemeDir,"history/")
        if not os.path.exists(cur_dir):
            os.makedirs(cur_dir)
        return cur_dir
    #同步进行
    def SearchTheme(self):
        url='https://www.xiaohongshu.com/'
        
        # exist_cookie=exist_cookies(url)
        
        # if exist_cookie:
        #     load_cookies(self.wp,url)

        self.wp.get(url)
        
        #等待登录
        if RedBookSearch.is_first:
            time.sleep(5)
        # time.sleep(.5 if exist_cookie else 5.0) 
        # if not exist_cookie:
        #     save_cookies(self.wp,url)
        
        
        search_input = self.wp.ele('xpath://input[@class="search-input"]')
        search_input.clear()
        search_input.input(f'{self.theme_name}\n')
        time.sleep(.4)

        
        seach_button=self.wp.ele('xpath://div[@class="search-icon"]')
        if not seach_button:
            sys.exit(0)
        self.wp.ele('xpath://div[@class="search-icon"]').click()

    def StopListen(self):
        self.wp.listen.stop()

    def CloseBrowser(self):
        self.wp.quit()

        
    async def SearchNotes(self):

        self.wp.listen.start(['web/v1/search/notes',"api/sns/web/v1/feed"]) 
        packet = self.wp.listen.wait()
        cur_dir=self.ThemeDir
        os.makedirs(cur_dir,exist_ok=True)


        sec_i=0
        
        secManager=SectionManager(self.wp)
        secManager.update()
        while (secManager.count())>0:
            
            if sec_i>self.search_count:
                self.stoped=True
                break
            # secManager.update()
            

            sec=next(secManager.next()) 
            if not sec:
                continue
            await asyncio.sleep(.1)
            
            try:
                sec.click()
            except:
                secManager.resume_cur()
                secManager.update()
                continue
            pack=self.wp.listen.wait()
            body=pack.response.body
            
            #非空才写入
            if body:
                sec_i+=1
                
                
                
                tasks=[
                    self.data_queue.put(body),
                #异步写入临时文件
                    read_write_async(json.dumps(body,indent=4,ensure_ascii=False),
                                    os.path.join(self.ThemeHistoryDir,f"{sec_i:04d}.json"), "w", encoding="utf-8")
                    
                ]
                await asyncio.gather(*tasks)




            close_flag=self.wp.ele('xpath://div[@class="close close-mask-dark"]')
            if close_flag:
                try:
                    close_flag.click()
                except:
                    continue

                
            
                

    async def HandleNotes(self):
        lst=[]
        
        while not (self.stoped and self.data_queue.empty()):
            if self.data_queue.empty():
                await asyncio.sleep(.1)
                continue
            
            data=await self.data_queue.get()
            if data:
                noteinfo= await ParseOrgNote(data)
                if noteinfo:
                    lst.append(noteinfo)
                    
                    noteinfo.set_root_dir(self.ThemeDir)
                    await noteinfo.write_to_notepad()
                self.cur_index+=1
            self.data_queue.task_done()
        if lst:
            df=pd.DataFrame([info.__dict__ for info in lst], columns=lst[0].__dict__.keys())
            if not df is None:
                df.sort_values(by=["thumbs","create_time"], ascending=[False, True],inplace=True)
            return df

    def __CreateNotes():
        
        index=0
        for root, dirs, files in os.walk(setting.redbook_history_dir):
            for name in files:
                if not name.endswith(".json"):
                    continue
                with open(os.path.join(root,name),"r",encoding="utf-8") as f:
                    content= f.read()
                    noteinfo= ParseOrgNote(content)
                    if noteinfo:
                        noteinfo.set_root_dir(setting.redbook_notes_dir,index)
                        noteinfo.write_to_notepad()
                    index+=1
                    
                    

    
async def to_theme_word(theme_name,root_dir,dict_data):
    await asyncio.sleep(.1)
        # 创建一个新的文档
    document = Document()
    #整理到word中
    for note_info in dict_data:
        info=NoteInfo(**note_info)
        info.to_word(document)
    document.save(os.path.join(root_dir,f"{theme_name}.docx"))
    
    
                    
class RedBookSearchs:
    def __init__(self,themes:list,root_dir:str=setting.redbook_notes_dir,search_count:int=5,queue_count:int=100) -> None:
        self.wp=WebPage()
        self.themes=themes
        self.root_dir=root_dir
        self.search_count=search_count
        self.queue_count=queue_count
        
    def InfoToExcel(self,root_dir, pds:list):
        outPath = os.path.join(root_dir,"notes_tab.xlsx")
        with pd.ExcelWriter(outPath) as writer:
            for theme,df in pds:
                if not df is None:
                    df.to_excel(writer, sheet_name=theme)
 
 
        
    #无法异步处理，浏览器只能有一个
    async def AsyncDumps(self):
        sem= asyncio.Semaphore( os.cpu_count()+4 )
        async def fetch_theme(theme):
            async with sem:
                search=RedBookSearch(None,theme,self.root_dir,self.search_count,self.queue_count)
                search.SearchTheme()
                
                tasks=[search.SearchNotes(),search.HandleNotes()]
                task_list=[asyncio.create_task(task) for task in tasks]
                results= await asyncio.gather(*task_list)
                search.StopListen()
                return results[1]
            
        
        tasks=[fetch_theme(theme) for theme in self.themes]
        results= await asyncio.gather(*tasks)
        
        dfs=[result for result in results if not result is None]
        self.InfoToExcel(self.root_dir,dfs)
        
    async def Dumps(self):
        
        dfs=[]
        root_dir=self.root_dir
        cache_dir=os.path.join(root_dir,"cache")
        os.makedirs(cache_dir,exist_ok=True)
        start_time=time.time()
        for index,theme in enumerate(self.themes) :
            cur_time= time.time()
            if index>0 and RedBookSearch.is_first:
                RedBookSearch.is_first=False
                
            search=RedBookSearch(self.wp,theme,root_dir,self.search_count,self.queue_count)
            search.SearchTheme()
            
            tasks=[search.SearchNotes(),search.HandleNotes()]
            task_list=[asyncio.create_task(task) for task in tasks]
            results= await asyncio.gather(*task_list)
            df:pd.DataFrame= results[1]
            if not df is None:
                dfs.append((theme,df) )
                
                #临时数据，缓存使用、
                dict_data=df.to_dict("records")
                json_data=json.dumps(dict_data,indent=4,ensure_ascii=False)
                
                tasks=[read_write_async(json_data, os.path.join(cache_dir,f"{theme}.json"),mode="w",encoding="utf-8"),
                       to_theme_word(theme_name=theme,root_dir=root_dir,dict_data=dict_data)
                       ]
                    
                await asyncio.gather(*tasks)

            
            logger.info(f"查询【{theme}】,用时{time.time()- cur_time}")
            

            search.StopListen()
            
        self.InfoToExcel(root_dir,dfs)
        logger.info(f"一共用时{time.time()- start_time}")
        
        

if __name__ == '__main__':
    wp=WebPage()
    # themes=["补气血吃什么","黄芪",'麦冬','怀山药']
    themes=["薏米"]
    
    redbook=RedBookSearchs(themes=themes,root_dir=setting.redbook_notes_dir,search_count=50,queue_count=100)
    asyncio.run(redbook.Dumps()) 
    
