import os
import re
from collections import Counter
import json
import pandas as pd
import argparse
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

import __init__
from base.com_log import logger as logger

def create_dir_recursive(path):
    try:
        os.makedirs(path, exist_ok=True)
        logger.trace(f"Directory created: {path}")
    except OSError as error:
        logger.error(f"Directory creation failed: {error}")
        
# def init_logger(dir_name :str):
#     create_dir_recursive(dir_name)
#     log_path=os.path.join(dir_name,"日志.log")
#     # 创建一个handler，用于写入日志文件
#     fh = logging.FileHandler(log_path)
#     fh.setLevel(logging.DEBUG)

#     # 再创建一个handler，用于输出到控制台
#     ch = logging.StreamHandler()
#     ch.setLevel(logging.DEBUG)

#     # 定义handler的输出格式
#     formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
#     fh.setFormatter(formatter)
#     ch.setFormatter(formatter)

#     # 给logger添加handler
#     logger.addHandler(fh)
#     logger.addHandler(ch)




class ThemeItem:
    def __init__(self, name:str, count:int):
        self.name = name
        self.count = count
    def __str__(self):
        return f"name: {self.name}, count: {self.count}"
    
    


class Heading:
    def __init__(self, num:str, name:str, theme:list):
        self.num = num
        self.name = name
        self.theme = theme
        
    def __str__(self):
        
        theme_list_str=[str(item) for item in self.theme]
        
        theme_str=";".join(theme_list_str)

        return f"num: {self.num}, name: {self.name}, theme:{theme_str}"

def parse_heading(s):
    res_list = []
    
    sec_pattern = "\n\n\n"
    head_pattern = r'(\d{2}_\d{3})\.\s+(\S+)'
    # theme_pattern = r'#([^#\s]*)[^\w#\]\)\】]+'
    theme_pattern =  r'#([^#\s#\]\)\】,，。\?\(\)!！）（]*)'
    split_all=s.split(sec_pattern)
    for oneHead in split_all:
        head_list=[]
        detail=oneHead.split('\n')
        if(len(detail)<2):
            continue
        
        
        matches=re.findall(head_pattern, detail[0])
        if(not matches ):
            continue
        
        
        for item in detail[1:]:
            if item.find('#')==-1:
                continue
            theme=re.findall(theme_pattern,item)
            if not theme:
                continue
            head_list.extend( [item.strip() for item in theme if item.strip()])
        if len(head_list)<1:
            continue
        theme_counter=Counter(head_list)
        
        match=matches[0]
        
        
        
        # theme_list=list(dict.fromkeys(head_list)) 
        theme_list=[ ThemeItem(key,theme_counter[key]) for key in theme_counter]  
        res_list.append(Heading(match[0],match[1],theme_list))
    return res_list


def handle_each(filePath: str):
    content=""
    with open(filePath, 'r', encoding='utf8') as file:
    # 读取文件内容
        content = file.read()
    fileName= os.path.splitext(os.path.basename(filePath))[0]
    
    
    arr=parse_heading(content)

    dest_map={}
    for  item in arr:
        for key in item.theme:
            themestr=key.name
            num=item.num
            if key.count>1:
                num+=f"({key.count})"
            if themestr not in dest_map:
                dest_map[themestr]=[num]
            else:
                dest_map[themestr].append(num)
    
    result_map=sorted(dest_map.items(), key=lambda x: len(x[1]), reverse=True)
    # print(result_map)  
    
    # dest_count={item[0]:len(item[1]) for item in result_map}
    # print(dest_count)    
    
    # # 将JSON格式的字符串写入文件
    # json_data=json.dumps(dest_count,ensure_ascii=False)
    # with open("E:/小红书竞标/修改后 - 副本 - 副本/01_故里有茶-薏米2_result_data.json", "w",encoding='utf8') as write_file:
    #     write_file.write(json_data)
    
    
    return (fileName,result_map) 


    
    
# 将 DataFrame 添加到 Word 文档中
def add_df_to_document(df: pd.DataFrame, doc: Document,titleName:str):
    p=doc.add_paragraph(titleName)
    p.alignment = 1  # 1 表示居中对齐
    doc.add_section(WD_SECTION.CONTINUOUS) 
    
    tbl = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = tbl.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name

    for _, row in df.iterrows():
        row_cells = tbl.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)
            
            
    # 自适应宽度 - 设置每一列的宽度
    for col in tbl.columns:
        for cell in col.cells:
            # 设置单元格宽度为Inches(1.5)，可以根据需要调整
            cell.width = Inches(1)

            
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 1  # 1 表示居中对齐
    # 换节符
    doc.add_section(WD_SECTION.NEW_PAGE)    
            
    # # 添加空行
    # p=doc.add_paragraph()
    # p.paragraph_format.space_after = Inches(1)  # 设置段后间距为12磅

def set_doc_font(doc: Document):
        # 设置默认字体为“宋体”
    default_font_name = 'SimSun'
    doc.styles['Normal'].font.name = default_font_name
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), default_font_name)
    # 遍历文档中的所有段落，设置字体
    for para in doc.paragraphs:
        para.style = doc.styles['Normal']

    # 遍历文档中的所有表格，设置字体
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles['Normal']
    
def Handle_all(dir_path:str):

    output_xl_path=os.path.join(dir_path,"统计结果.xlsx")
    # 创建一个ExcelWriter对象
    writer = pd.ExcelWriter(output_xl_path)
    
    # log_path=os.path.join(dir_path,"py.log")
    # # 创建一个handler，用于写入日志文件
    # fh = logging.FileHandler(log_path)
    # fh.setLevel(logging.DEBUG)

    # # 再创建一个handler，用于输出到控制台
    # ch = logging.StreamHandler()
    # ch.setLevel(logging.DEBUG)

    # # 定义handler的输出格式
    # formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    # fh.setFormatter(formatter)
    # ch.setFormatter(formatter)

    # # 给logger添加handler
    # logger.addHandler(fh)
    # logger.addHandler(ch)
    
    
    
    # 创建一个 Word 文档
    output_doc_path=os.path.join(dir_path,"统计结果.docx")
    doc = Document()

    
    for filename in os.listdir(dir_path):
        # 筛选出.txt文件
        if filename.endswith('.txt'):
            cur_file_path=os.path.join(dir_path,filename)
            (name,data)=handle_each(cur_file_path)
            if len(data)<1 or len(name)<1:
                logger.warning(f"解析出错：{cur_file_path}")    
                continue
            logger.trace(f"{cur_file_path}解析成功：\n{json.dumps(data,ensure_ascii=False)}")
            
            themeTitile="话题"
            df = pd.DataFrame(data,columns=[themeTitile,"关联日记"])
            df["数量"]=df["关联日记"].apply(len)
            column_order=[themeTitile,'数量','关联日记']
            df=df[column_order]
            
            df.to_excel(writer, sheet_name=name, index=False,)
            
            df.drop('关联日记', axis=1, inplace=True)
            
            # 添加第一个 DataFrame
            add_df_to_document(df, doc,name)
    # 保存Excel文件
    writer.close()
    
    
    # 记录一条日志
    logger.info(f"统计结果已保存到：{output_xl_path}")      
    set_doc_font(doc)
    doc.save(output_doc_path)
    logger.info(f"统计结果已保存到：{output_doc_path}")    
    


    
    

if __name__ == '__main__':
    logger.info("程序执行开始！")
    parser = argparse.ArgumentParser(description='分析话题信息(*.txt)')
    parser.add_argument('-i', '--input', type=str,  help='输入文件夹路径')
    args = parser.parse_args()
    
    input_agrs = args.input
    # if not input_agrs is None :
    #     input_agrs=input_agrs.replace("\\","/")
    if input_agrs is None or  not os.path.exists(input_agrs):
        warn_str=f"path {input_agrs} is None" if input_agrs is None else f"path {input_agrs} not exists"
        
        logger.error(f"参数不匹配，请检查输入路径！{warn_str}")
        sys.exit(0)

    Handle_all(input_agrs)
    logger.info("程序执行完毕！")
    
# 打包说明：pyinstaller --onefile --distpath exe -p . --distpath .\exe ThemeCount.py
#pyinstaller  --onefile --distpath exe -p . -p base -p integer -p red_book --add-data "config/settings.yaml:./config" --add-data "config/.secrets.yaml:./config" --distpath .\exe red_book/ThemeCount.py    


# 生成.\exe\ThemeCount.exe，运行即可

# exe运行
# ThemeCount.exe -i "E:/小红书竞标/参考话题/output"
 